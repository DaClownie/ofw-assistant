# app/utils/memo.py
import json
from pathlib import Path
from docx import Document

MEMO_PATH = Path("data/memo_store.json")
MEMO_PATH.parent.mkdir(parents=True, exist_ok=True)

def load_memos():
    if MEMO_PATH.exists():
        with open(MEMO_PATH, "r") as f:
            return json.load(f)
    return {}

def save_memo(case_id: str, memo_text: str, sources=None, memo_name=None):
    memos = load_memos()
    
    # If there's already a memo for this case, create a list
    if case_id in memos:
        # Convert single memo to list format if needed
        if not isinstance(memos[case_id], list):
            old_memo = memos[case_id]
            memos[case_id] = [old_memo]
        
        # Add new memo to the list
        memos[case_id].append({
            "memo_name": memo_name or f"Memo {len(memos[case_id]) + 1}",
            "memo_text": memo_text,
            "sources": sources or []
        })
    else:
        # First memo for this case
        memos[case_id] = [{
            "memo_name": memo_name or "Memo 1",
            "memo_text": memo_text,
            "sources": sources or []
        }]
    
    with open(MEMO_PATH, "w") as f:
        json.dump(memos, f, indent=2)

def export_memo_docx(case_id: str, memo_text: str, output_path: Path):
    doc = Document()
    doc.add_heading(f"Memo for Case: {case_id}", level=1)
    doc.add_paragraph(memo_text)
    output_path = output_path.resolve()
    print(f"Saving .docx to: {output_path}")
    doc.save(str(output_path))

from app.utils.vectorstore import get_chunks_by_file as vectordb_get_chunks_by_file
from app.utils.model_llama import tag_with_llama

def get_chunks_by_file(file_path: str):
    return vectordb_get_chunks_by_file(file_path)

def summarize_chunks(chunks):
    if not chunks:
        print("⚠️ summarize_chunks: No chunks received.")
        return "No content available"
    
    print(f"✅ summarize_chunks: Processing {len(chunks)} chunks.")
    
    # Combine all text
    full_text = " ".join([
        chunk.page_content if hasattr(chunk, "page_content") else str(chunk) 
        for chunk in chunks
    ])
    
    # Get tags for the full document
    from app.utils.controlled_smart_tagger import controlled_smart_tagger
    analysis = controlled_smart_tagger.analyze_text_comprehensive(full_text[:3000])
    
    # Create a structured summary
    summary = f"""
CONTENT: {full_text[:500]}{'...' if len(full_text) > 500 else ''}

IDENTIFIED ISSUES:
{chr(10).join(f"- {tag.replace('_', ' ').title()}" for tag in analysis['tags']) if analysis['tags'] else '- None'}

PRIORITY CONCERNS:
{chr(10).join(f"- {tag.replace('_', ' ').title()}" for tag in analysis['priority_tags']) if analysis['priority_tags'] else '- None'}
"""
    
    return summary.strip()