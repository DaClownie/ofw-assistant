# app/utils/unified_processor.py
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from langchain_core.documents import Document
import mimetypes
import shutil

from .parser import load_and_split_pdf
from .image import extract_text_from_image
from .audio import transcribe_audio
from .tagging import should_use_gpt4
from .model_gpt4 import tag_with_gpt4
from .model_llama import tag_with_llama
from .vectorstore import persist_chunks
from .memory import update_memory

class FileProcessor:
    """Unified file processor that handles all media types consistently"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.docx': self._process_document,
            '.eml': self._process_email,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image,
            '.heic': self._process_image,
            '.tiff': self._process_image,
            '.gif': self._process_image,
            '.bmp': self._process_image,
            '.mp3': self._process_audio,
            '.m4a': self._process_audio,
            '.mp4': self._process_video,
            '.mov': self._process_video,
            '.mkv': self._process_video,
            '.avi': self._process_video
        }
    
    def process_file(self, file_path: str, case_id: str) -> Dict:
        """Main entry point for processing any file type"""
        path_obj = Path(file_path)
        extension = path_obj.suffix.lower()
        
        if extension not in self.supported_types:
            raise ValueError(f"Unsupported file type: {extension}")
        
        # Create case-specific directory and move file
        case_dir = Path("data/case_files") / case_id.replace(" ", "_")
        case_dir.mkdir(parents=True, exist_ok=True)
        
        # Move file to case directory
        new_file_path = case_dir / path_obj.name
        if Path(file_path).resolve() != new_file_path.resolve():
            shutil.move(str(file_path), str(new_file_path))
            file_path = str(new_file_path)
        
        # Process the file and get structured results
        result = self.supported_types[extension](file_path)
        
        # Store in vector database
        if result['documents']:
            persist_chunks(result['documents'], doc_id=path_obj.name)
        
        # Update memory with consistent structure
        update_memory(
            file_path,
            tags=result['tags'],
            case_id=case_id,
            flags=result['flags'],
            transcript=result.get('transcript'),
            extra_meta={"path": file_path, "type": extension}
        )
        
        return result
    
    def _save_transcript_file(self, file_path: str, transcript: str):
        """Save transcript as a .txt file alongside the original"""
        if transcript and transcript.strip() and not transcript.startswith("[Error"):
            transcript_path = f"{file_path}.txt"
            try:
                with open(transcript_path, 'w', encoding='utf-8') as f:
                    f.write(transcript)
                print(f"Transcript saved to: {transcript_path}")
            except Exception as e:
                print(f"Failed to save transcript: {e}")
    
    def _extract_and_analyze_text(self, text: str, source_path: str) -> Dict:
        """Common text analysis pipeline"""
        if not text.strip():
            return {
                'documents': [],
                'tags': [],
                'flags': [],
                'transcript': text
            }
        
        # Create document for vector storage
        doc = Document(
            page_content=text, 
            metadata={"source": str(Path(source_path).resolve())}
        )
        
        # Import here to avoid circular imports
        from .controlled_smart_tagger import controlled_smart_tagger

        # Analyze content with controlled taxonomy only
        tags = controlled_smart_tagger.tag_text(text)
        flags = []  # No separate flagging - tags handle everything
        
        return {
            'documents': [doc],
            'tags': tags,
            'flags': flags,
            'transcript': text
        }
    
    def _process_pdf(self, file_path: str) -> Dict:
        """Process PDF files"""
        chunks = load_and_split_pdf(file_path)
        all_tags = set()
        
        # Import here to avoid circular imports
        from .controlled_smart_tagger import controlled_smart_tagger
        
        for chunk in chunks:
            chunk.metadata["source"] = str(Path(file_path).resolve())
            text = chunk.page_content
            
            # Analyze each chunk with controlled taxonomy
            tags = controlled_smart_tagger.tag_text(text)
            
            all_tags.update(tags)
        
        return {
            'documents': chunks,
            'tags': list(all_tags),
            'flags': [],
            'transcript': None
        }
    
    def _process_image(self, file_path: str) -> Dict:
        """Process image files with OCR"""
        text = extract_text_from_image(file_path)
        return self._extract_and_analyze_text(text, file_path)
    
    def _process_audio(self, file_path: str) -> Dict:
        """Process audio files with transcription"""
        transcript = transcribe_audio(file_path)
        self._save_transcript_file(file_path, transcript)
        return self._extract_and_analyze_text(transcript, file_path)
    
    def _process_video(self, file_path: str) -> Dict:
        """Process video files (extract audio for transcription)"""
        # Video processing uses same audio transcription
        transcript = transcribe_audio(file_path)
        self._save_transcript_file(file_path, transcript)
        return self._extract_and_analyze_text(transcript, file_path)
    
    def _process_text(self, file_path: str) -> Dict:
        """Process plain text files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return self._extract_and_analyze_text(text, file_path)
    
    def _process_document(self, file_path: str) -> Dict:
        """Process DOCX files"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return self._extract_and_analyze_text(text, file_path)
        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return self._extract_and_analyze_text("", file_path)
    
    def _process_email(self, file_path: str) -> Dict:
        """Process email files"""
        try:
            import email
            from email import policy
            
            with open(file_path, 'rb') as f:
                msg = email.message_from_binary_file(f, policy=policy.default)
            
            # Extract email content
            subject = msg.get('subject', '')
            sender = msg.get('from', '')
            date = msg.get('date', '')
            
            # Get email body
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body += part.get_payload(decode=True).decode('utf-8', errors='ignore')
            else:
                body = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            
            # Combine into full text
            text = f"Subject: {subject}\nFrom: {sender}\nDate: {date}\n\n{body}"
            
            return self._extract_and_analyze_text(text, file_path)
        except Exception as e:
            print(f"Error processing email: {e}")
            return self._extract_and_analyze_text("", file_path)